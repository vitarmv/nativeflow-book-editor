import streamlit as st
from docx import Document
from docx.shared import Inches, Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, ns 
import google.generativeai as genai
from io import BytesIO
import time
import os
import re
import uuid
import itertools

# --- LIBRERÍAS PRO ---
import pyphen  
import mammoth 
from bs4 import BeautifulSoup 
from ebooklib import epub 

# --- 1. CONFIGURACIÓN GLOBAL ---
st.set_page_config(page_title="Suite Autores 360 ULTIMATE", page_icon="📚", layout="wide")

st.markdown("""
<style>
    .stProgress > div > div > div > div { background-color: #4CAF50; }
    .block-container { padding-top: 2rem; }
    div[data-testid="stSidebar"] { background-color: #f8f9fa; }
    h1 { color: #2c3e50; }
</style>
""", unsafe_allow_html=True)

# --- 2. DICCIONARIO DE TEMAS ---
THEMES = {
    "Neutro (Estándar)": {"font": "Calibri", "header": "Calibri", "size": 11},
    "Romance / Fantasía (Serif)": {"font": "Garamond", "header": "Garamond", "size": 12},
    "Thriller / Crimen (Sharp)": {"font": "Georgia", "header": "Arial Black", "size": 11},
    "No Ficción / Negocios": {"font": "Arial", "header": "Arial", "size": 10}
}

# --- 3. BARRA LATERAL ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/3145/3145765.png", width=80)
    st.title("Centro de Mando")
    
    try:
        api_key = st.secrets["GOOGLE_API_KEY"]
        genai.configure(api_key=api_key)
        st.success("✅ Motor IA Activo")
    except:
        st.error("❌ Falta API Key en Secrets")
        st.stop()
        
    st.divider()
    
    selected_module = st.radio(
        "Selecciona Herramienta:",
        [
            "1. 💎 Auditor & Corrector IA",
            "2. 📏 Maquetador KDP PRO (Papel)",
            "3. 📲 Workbook Cleaner (Kindle)",
            "4. 🧼 Limpiador Rápido",
            "5. ⚡ Generador EPUB (eBook)"
        ]
    )
    
    st.divider()
    MODEL_NAME = 'models/gemini-flash-latest' 
    model = genai.GenerativeModel(MODEL_NAME)

# --- 4. FUNCIONES AUXILIARES ---

def create_element(name): return OxmlElement(name)
def create_attribute(element, name, value): element.set(ns.qn(name), value)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_run = paragraph.add_run()
    t1 = create_element('w:fldChar'); create_attribute(t1, 'w:fldCharType', 'begin'); page_run._r.append(t1)
    t2 = create_element('w:instrText'); create_attribute(t2, 'xml:space', 'preserve'); t2.text = "PAGE"; page_run._r.append(t2)
    t3 = create_element('w:fldChar'); create_attribute(t3, 'w:fldCharType', 'end'); page_run._r.append(t3)

def enable_native_hyphenation(doc):
    settings = doc.settings.element
    hyphenation_zone = OxmlElement('w:autoHyphenation')
    create_attribute(hyphenation_zone, 'w:val', 'true')
    settings.append(hyphenation_zone)

def prevent_runts_in_paragraph(paragraph):
    text = paragraph.text.strip()
    if not text or len(text) < 20: return 
    last_space = text.rfind(' ')
    if last_space != -1:
        paragraph.text = text[:last_space] + "\u00A0" + text[last_space+1:]

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def stitch_paragraphs(doc):
    for i in range(len(doc.paragraphs) - 2, -1, -1):
        p_curr = doc.paragraphs[i]
        p_next = doc.paragraphs[i+1]
        text_curr = p_curr.text.strip()
        text_next = p_next.text.strip()
        if not text_curr or not text_next: continue
        if p_curr.style.name.startswith('Heading') or p_next.style.name.startswith('Heading'): continue
        if text_curr[-1] not in ['.', '!', '?', '"', '”', ':']:
            p_curr.text = text_curr + " " + text_next
            delete_paragraph(p_next)

def nuclear_clean(text):
    if not text: return text
    text = text.replace('\n', ' ').replace('\r', ' ').replace('\v', ' ').replace('\f', ' ')
    return " ".join(text.split())

def clean_markdown(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text) 
    text = re.sub(r'\*(.*?)\*', r'\1', text)      
    text = re.sub(r'__(.*?)__', r'\1', text)      
    text = re.sub(r'^#+\s*', '', text) 
    return nuclear_clean(text).strip()

def call_api(prompt, temp=0.7):
    for _ in range(3):
        try: return model.generate_content(prompt, generation_config={"temperature": temp}).text.strip()
        except: time.sleep(1)
    return "[ERROR API]"

# ==============================================================================
# MÓDULO 1: AUDITOR & CORRECTOR
# ==============================================================================
if "1." in selected_module:
    st.header("💎 Auditoría & Corrección IA")
    uploaded_file = st.file_uploader("Sube tu manuscrito", type=["docx"], key="mod1")
    if uploaded_file:
        doc = Document(uploaded_file)
        tab1, tab2 = st.tabs(["📊 Auditoría de Calidad", "🚀 Corrección de Estilo"])
        with tab1:
            if st.button("🔍 Iniciar Auditoría"):
                audit_doc = Document()
                audit_doc.add_heading("Reporte de Auditoría", 0)
                p_bar = st.progress(0)
                for i, p in enumerate(doc.paragraphs):
                    if len(p.text) > 15:
                        res = call_api(f"Analyze the following text for grammar or flow issues. Output 'CLEAN' if perfect, or describe the issue. Text: '{p.text[:400]}'")
                        if "CLEAN" not in res: audit_doc.add_paragraph(f"Párrafo {i+1}: {res}")
                    p_bar.progress((i+1)/len(doc.paragraphs))
                bio = BytesIO(); audit_doc.save(bio)
                st.download_button("⬇️ Descargar Reporte", bio.getvalue(), "Auditoria.docx")
        with tab2:
            if st.button("🚀 Re-escribir con IA"):
                new_doc = Document()
                p_bar = st.progress(0)
                for i, p in enumerate(doc.paragraphs):
                    if len(p.text) > 5:
                        res = call_api(f"Improve the flow and style of this text, keep original meaning: '{p.text}'")
                        new_doc.add_paragraph(clean_markdown(res))
                    else: new_doc.add_paragraph("")
                    p_bar.progress((i+1)/len(doc.paragraphs))
                bio = BytesIO(); new_doc.save(bio)
                st.download_button("⬇️ Descargar Corregido", bio.getvalue(), "Manuscrito_IA.docx")

# ==============================================================================
# MÓDULO 2: MAQUETADOR KDP PRO (PAPEL)
# ==============================================================================
elif "2." in selected_module:
    st.header("📏 Maquetador KDP PRO (Papel)")
    col1, col2 = st.columns(2)
    with col1:
        size = st.selectbox("Tamaño de impresión:", ["6 x 9 pulgadas", "5 x 8 pulgadas", "8.5 x 11 pulgadas"])
        theme_choice = st.selectbox("🎨 Tema Visual:", list(THEMES.keys())) 
    with col2:
        margins = st.radio("Configuración de Márgenes:", ["Espejo (Doble Cara)", "Normales"])
    
    st.markdown("---")
    col3, col4 = st.columns(2)
    with col3:
        fix_titles = st.checkbox("📎 Forzar Títulos en Hoja Nueva", value=True)
        pro_start = st.checkbox("✨ Activar Letra Capital al Inicio", value=True)
    with col4:
        reconstruct = st.checkbox("🔗 Unir párrafos rotos (Reconstructor)", value=True)
        justify_text = st.checkbox("📄 Justificar + Silabeo", value=True)
        add_numbers = st.checkbox("🔢 Agregar Números de Página", value=True)
        fix_runts = st.checkbox("🛡️ Evitar palabras sueltas (Runts)", value=True)

    uploaded_file = st.file_uploader("Sube manuscrito (.docx)", type=["docx"], key="mod2")
    if uploaded_file and st.button("🛠️ Procesar Libro"):
        doc = Document(uploaded_file)
        theme = THEMES[theme_choice]
        
        if reconstruct: stitch_paragraphs(doc)
        if justify_text: 
            try: enable_native_hyphenation(doc) 
            except: pass

        if "6 x 9" in size: w, h = Inches(6), Inches(9)
        elif "5 x 8" in size: w, h = Inches(5), Inches(8)
        else: w, h = Inches(8.5), Inches(11)
        
        for section in doc.sections:
            section.page_width = w; section.page_height = h
            section.top_margin = Inches(0.75); section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.8); section.right_margin = Inches(0.6)
            if "Espejo" in margins: section.mirror_margins = True; section.gutter = Inches(0.15)
            if add_numbers:
                p_foot = section.footer.paragraphs[0]; p_foot.text = ""
                add_page_number(p_foot)
                p_foot.style.font.name = theme['font']; p_foot.style.font.size = Pt(10)

        style = doc.styles['Normal']
        style.font.name = theme['font']; style.font.size = Pt(theme['size'])
        style.paragraph_format.line_spacing = 1.25; style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.widow_control = True
        if justify_text: style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for h in ['Heading 1', 'Heading 2']:
            try:
                h_style = doc.styles[h]
                h_style.font.name = theme['header']; h_style.font.color.rgb = RGBColor(0,0,0)
                h_style.paragraph_format.space_before = Pt(0); h_style.paragraph_format.space_after = Pt(30)
                h_style.alignment = WD_ALIGN_PARAGRAPH.CENTER; h_style.paragraph_format.page_break_before = True
            except: pass

        total_p = len(doc.paragraphs)
        p_bar = st.progress(0)
        previous_was_heading = False 
        
        for i, p in enumerate(doc.paragraphs):
            text_clean = p.text.strip()
            if len(text_clean) < 2: continue 
            is_h = p.style.name.startswith('Heading') or (len(text_clean) < 60 and (re.match(r'^(chapter|cap[íi]tulo)\b', text_clean, re.I) or text_clean.isupper()))
            
            if is_h:
                previous_was_heading = True
                p.style = doc.styles['Heading 1']
                p.text = "\n" + text_clean.upper() 
                if fix_titles: p.paragraph_format.keep_with_next = True
            else:
                if fix_runts and len(text_clean) > 50: prevent_runts_in_paragraph(p)
                if justify_text: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if pro_start and previous_was_heading:
                    if len(text_clean) > 1:
                        char = text_clean[0]; rest = text_clean[1:]
                        p.text = ""; run = p.add_run(char)
                        run.font.name = theme['header']; run.font.size = Pt(theme['size']+5); run.bold = True
                        p.add_run(rest).font.name = theme['font']
                    previous_was_heading = False
                else: previous_was_heading = False
            if i % 10 == 0: p_bar.progress((i+1)/total_p)
        
        bio = BytesIO(); doc.save(bio)
        st.success("✅ Maquetación de papel terminada.")
        st.download_button("⬇️ Descargar DOCX", bio.getvalue(), "Maquetado_Papel.docx")

# ==============================================================================
# MÓDULO 3: WORKBOOK CLEANER
# ==============================================================================
elif "3." in selected_module:
    st.header("📲 Workbook Cleaner (Kindle)")
    cta_text = st.text_area("Texto CTA:", "🛑 (Ejercicio): Completa esto en tu Cuaderno.", height=80)
    uploaded_file = st.file_uploader("Sube manuscrito", key="mod3")
    if uploaded_file and st.button("Limpiar"):
        doc = Document(uploaded_file)
        for p in doc.paragraphs:
            if re.search(r"([_.\-]){4,}", p.text): p.text = cta_text 
        bio = BytesIO(); doc.save(bio)
        st.download_button("⬇️ Descargar", bio.getvalue(), "Ebook.docx")

# ==============================================================================
# MÓDULO 4: LIMPIADOR
# ==============================================================================
elif "4." in selected_module:
    st.header("☢️ Limpiador 'Nuclear'")
    uploaded_file = st.file_uploader("Sube docx", key="mod4")
    if uploaded_file and st.button("Limpiar"):
        doc = Document(uploaded_file)
        for p in doc.paragraphs:
            if p.text: p.text = nuclear_clean(p.text)
        bio = BytesIO(); doc.save(bio)
        st.download_button("⬇️ Descargar", bio.getvalue(), "Limpio.docx")

# ==============================================================================
# MÓDULO 5: GENERADOR EPUB (V6.6 - LETRA CAPITAL FORZADA)
# ==============================================================================
elif "5." in selected_module:
    st.header("⚡ Generador EPUB 6.6 (Drop Cap Force)")
    uploaded_file = st.file_uploader("Sube Manuscrito (DOCX procesado)", key="mod5")
    
    col1, col2, col3 = st.columns(3)
    with col1: title = st.text_input("Título", "Mi Libro")
    with col2: author = st.text_input("Autor", "Autor")
    with col3: lang = st.selectbox("Idioma", ["es", "en", "fr", "it", "pt"]) 
    
    if uploaded_file and st.button("Convertir"):
        # 1. LIMPIEZA PREVIA
        doc_temp = Document(uploaded_file)
        # Elimina Enters de títulos
        for p in doc_temp.paragraphs:
            if p.style.name.startswith('Heading'):
                p.text = p.text.replace('\n', '').strip()

        # Eliminar párrafos vacíos "basura"
        paragraphs_to_delete = []
        for i in range(len(doc_temp.paragraphs) - 1):
            curr = doc_temp.paragraphs[i]
            next_p = doc_temp.paragraphs[i+1]
            if curr.style.name.startswith('Heading') and not next_p.text.strip():
                paragraphs_to_delete.append(next_p)
        for p in paragraphs_to_delete:
            p._element.getparent().remove(p._element)

        buf = BytesIO(); doc_temp.save(buf); buf.seek(0)
        
        # 2. CONFIGURACIÓN
        book = epub.EpubBook()
        book.set_identifier(str(uuid.uuid4()))
        book.set_title(title); book.set_language(lang); book.add_author(author)
        
        # 3. MAPA DE ESTILOS (INDICE)
        style_map = """
        p[style-name^='Heading'] => h1:fresh
        p[style-name^='Título'] => h1:fresh
        """
        result = mammoth.convert_to_html(buf, style_map=style_map)
        
        # 4. INYECTOR DE ETIQUETA FÍSICA (LA SOLUCIÓN REAL)
        soup = BeautifulSoup(result.value, 'html.parser')
        
        for h1 in soup.find_all('h1'):
            next_p = h1.find_next_sibling()
            # Buscar el siguiente párrafo real
            while next_p and (next_p.name != 'p' or not next_p.get_text().strip()):
                next_p = next_p.find_next_sibling()
            
            if next_p and next_p.name == 'p':
                text = next_p.get_text()
                if len(text) > 1:
                    # Envolver la primera letra en un SPAN con clase
                    first_char = text[0]
                    rest = text[1:]
                    new_html = f'<span class="dropcap">{first_char}</span>{rest}'
                    # Reemplazar el contenido del párrafo
                    new_tag = BeautifulSoup(new_html, 'html.parser')
                    next_p.clear()
                    next_p.append(new_tag)

        # 5. CSS BLINDADO (Apunta a la clase .dropcap)
        css = """<style>
            h1 { 
                margin-top: 2em; 
                text-align: center; 
                color: black;
                margin-bottom: 1em;
            }
            p { text-align: justify; text-indent: 1em; line-height: 1.5em; margin-top: 0; }
            
            /* CLASE dropcap FÍSICA */
            span.dropcap {
                float: left;
                font-size: 3.5em !important;
                font-weight: bold;
                line-height: 0.8em;
                margin-right: 0.1em;
                margin-top: -0.1em;
                color: black;
            }
        </style>"""

        content = soup.body if soup.body else soup
        chapters = []
        headers = soup.find_all('h1')
        
        if not headers:
            c = epub.EpubHtml(title="Inicio", file_name="chap_1.xhtml", lang=lang)
            c.content = css + str(content)
            book.add_item(c); chapters.append(c)
        else:
            curr_h, curr_t, count = "", "Inicio", 0
            for elem in content.children:
                if elem.name == 'h1':
                    if curr_h.strip():
                        count += 1
                        c = epub.EpubHtml(title=curr_t, file_name=f"c_{count}.xhtml", lang=lang)
                        page_break = '<div style="page-break-before:always;"></div>' if count > 1 else ""
                        c.content = css + page_break + f"<h1>{curr_t}</h1>{curr_h}"
                        book.add_item(c); chapters.append(c)
                    curr_t, curr_h = elem.get_text(), ""
                else: curr_h += str(elem)
            
            if curr_h.strip():
                count += 1
                c = epub.EpubHtml(title=curr_t, file_name=f"c_{count}.xhtml", lang=lang)
                page_break = '<div style="page-break-before:always;"></div>' if count > 1 else ""
                c.content = css + page_break + f"<h1>{curr_t}</h1>{curr_h}"
                book.add_item(c); chapters.append(c)

        book.toc = tuple(chapters)
        book.add_item(epub.EpubNcx()); book.add_item(epub.EpubNav())
        book.spine = ['nav'] + chapters
        
        bio_ep = BytesIO(); epub.write_epub(bio_ep, book, {})
        st.success("✅ EPUB generado: Idioma + Índice + Letra Grande Forzada.")
        st.download_button("⬇️ Descargar EPUB", bio_ep.getvalue(), f"{title}.epub")
